"""
File Upload Router - Handle file uploads for agent conversations
Supports PDF, CSV, Excel, DOCX, TXT, and image files
"""

import base64
import io
import json
import logging
import os
import uuid
from datetime import datetime
from typing import Dict, List, Optional, Any

from fastapi import APIRouter, File, HTTPException, UploadFile, Depends, Form
from pydantic import BaseModel

from src.database import get_db_connection
from src.dependencies import get_current_user_id
from uuid import UUID

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/files", tags=["files"])

# Supported file types and their MIME types
SUPPORTED_TYPES = {
    # Documents
    "application/pdf": "pdf",
    "application/msword": "doc",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/markdown": "md",
    "text/csv": "csv",
    "application/vnd.ms-excel": "xls",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/json": "json",
    # Images
    "image/png": "png",
    "image/jpeg": "jpg",
    "image/gif": "gif",
    "image/webp": "webp",
}

MAX_FILE_SIZE = 25 * 1024 * 1024  # 25MB


class FileUploadResponse(BaseModel):
    id: str
    filename: str
    content_type: str
    size: int
    extracted_text: Optional[str] = None
    data_uri: Optional[str] = None  # For images
    metadata: Dict[str, Any] = {}


class ExtractedContent(BaseModel):
    text: str
    tables: List[Dict[str, Any]] = []
    metadata: Dict[str, Any] = {}


def extract_text_from_pdf(file_bytes: bytes) -> ExtractedContent:
    """Extract text from PDF files"""
    try:
        import PyPDF2
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        
        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
        
        return ExtractedContent(
            text="\n\n".join(text_parts),
            metadata={"pages": len(pdf_reader.pages)}
        )
    except ImportError:
        # Fallback without PyPDF2
        logger.warning("PyPDF2 not installed, cannot extract PDF text")
        return ExtractedContent(text="[PDF text extraction requires PyPDF2]")
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        return ExtractedContent(text=f"[Error extracting PDF: {str(e)}]")


def extract_text_from_docx(file_bytes: bytes) -> ExtractedContent:
    """Extract text from DOCX files"""
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract from tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append({"rows": table_data})
        
        return ExtractedContent(
            text="\n\n".join(paragraphs),
            tables=tables,
            metadata={"paragraphs": len(paragraphs), "tables": len(tables)}
        )
    except ImportError:
        logger.warning("python-docx not installed, cannot extract DOCX text")
        return ExtractedContent(text="[DOCX text extraction requires python-docx]")
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        return ExtractedContent(text=f"[Error extracting DOCX: {str(e)}]")


def extract_text_from_csv(file_bytes: bytes) -> ExtractedContent:
    """Extract data from CSV files"""
    try:
        import csv
        
        # Try to detect encoding
        content = file_bytes.decode('utf-8', errors='replace')
        reader = csv.reader(io.StringIO(content))
        
        rows = list(reader)
        if not rows:
            return ExtractedContent(text="[Empty CSV file]")
        
        # Format as text
        headers = rows[0] if rows else []
        text_parts = [", ".join(headers)]
        
        for row in rows[1:]:
            text_parts.append(", ".join(row))
        
        return ExtractedContent(
            text="\n".join(text_parts),
            tables=[{"headers": headers, "rows": rows[1:]}],
            metadata={"rows": len(rows), "columns": len(headers)}
        )
    except Exception as e:
        logger.error(f"Error extracting CSV: {e}")
        return ExtractedContent(text=f"[Error extracting CSV: {str(e)}]")


def extract_text_from_excel(file_bytes: bytes) -> ExtractedContent:
    """Extract data from Excel files"""
    try:
        import openpyxl
        
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        text_parts = []
        tables = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            
            rows = []
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                rows.append(row_values)
                text_parts.append(", ".join(row_values))
            
            if rows:
                tables.append({
                    "sheet": sheet_name,
                    "headers": rows[0] if rows else [],
                    "rows": rows[1:] if len(rows) > 1 else []
                })
        
        return ExtractedContent(
            text="\n".join(text_parts),
            tables=tables,
            metadata={"sheets": len(wb.sheetnames)}
        )
    except ImportError:
        logger.warning("openpyxl not installed, cannot extract Excel text")
        return ExtractedContent(text="[Excel extraction requires openpyxl]")
    except Exception as e:
        logger.error(f"Error extracting Excel: {e}")
        return ExtractedContent(text=f"[Error extracting Excel: {str(e)}]")


def extract_text_from_txt(file_bytes: bytes) -> ExtractedContent:
    """Extract text from plain text files"""
    try:
        # Try UTF-8 first, then fallback
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text = file_bytes.decode('latin-1')
        
        return ExtractedContent(
            text=text,
            metadata={"characters": len(text), "lines": text.count('\n') + 1}
        )
    except Exception as e:
        logger.error(f"Error reading text file: {e}")
        return ExtractedContent(text=f"[Error reading text: {str(e)}]")


def extract_text_from_json(file_bytes: bytes) -> ExtractedContent:
    """Extract content from JSON files"""
    try:
        data = json.loads(file_bytes.decode('utf-8'))
        formatted = json.dumps(data, indent=2)
        
        return ExtractedContent(
            text=formatted,
            metadata={"type": type(data).__name__}
        )
    except Exception as e:
        logger.error(f"Error parsing JSON: {e}")
        return ExtractedContent(text=f"[Error parsing JSON: {str(e)}]")


def process_file(file_bytes: bytes, content_type: str, filename: str) -> ExtractedContent:
    """Process a file and extract its content"""
    file_ext = SUPPORTED_TYPES.get(content_type, "").lower()
    
    if content_type == "application/pdf":
        return extract_text_from_pdf(file_bytes)
    elif content_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        return extract_text_from_docx(file_bytes)
    elif content_type == "text/csv":
        return extract_text_from_csv(file_bytes)
    elif content_type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
        return extract_text_from_excel(file_bytes)
    elif content_type in ["text/plain", "text/markdown"]:
        return extract_text_from_txt(file_bytes)
    elif content_type == "application/json":
        return extract_text_from_json(file_bytes)
    elif content_type.startswith("image/"):
        # For images, we return a data URI that can be sent to vision models
        b64_data = base64.b64encode(file_bytes).decode('utf-8')
        data_uri = f"data:{content_type};base64,{b64_data}"
        return ExtractedContent(
            text=f"[Image: {filename}]",
            metadata={"data_uri": data_uri, "is_image": True}
        )
    else:
        return ExtractedContent(text=f"[Unsupported file type: {content_type}]")


@router.post("/upload", response_model=FileUploadResponse)
async def upload_file(
    file: UploadFile = File(...),
    user_id: UUID = Depends(get_current_user_id)
):
    """
    Upload a file and extract its content.
    Supports: PDF, DOCX, DOC, TXT, CSV, XLS, XLSX, JSON, and images.
    Returns extracted text that can be used in agent conversations.
    """
    # Validate file type
    content_type = file.content_type or "application/octet-stream"
    
    if content_type not in SUPPORTED_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {content_type}. Supported types: {list(SUPPORTED_TYPES.values())}"
        )
    
    # Read file
    file_bytes = await file.read()
    
    # Check file size
    if len(file_bytes) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size is {MAX_FILE_SIZE / 1024 / 1024}MB"
        )
    
    # Generate file ID
    file_id = str(uuid.uuid4())
    filename = file.filename or f"upload_{file_id}"
    
    # Extract content
    extracted = process_file(file_bytes, content_type, filename)
    
    # Prepare response
    response = FileUploadResponse(
        id=file_id,
        filename=filename,
        content_type=content_type,
        size=len(file_bytes),
        extracted_text=extracted.text if not extracted.metadata.get("is_image") else None,
        data_uri=extracted.metadata.get("data_uri"),
        metadata={
            **extracted.metadata,
            "tables": extracted.tables if extracted.tables else None,
            "uploaded_at": datetime.utcnow().isoformat()
        }
    )
    
    logger.info(f"File uploaded: {filename} ({content_type}, {len(file_bytes)} bytes)")
    
    return response


@router.post("/upload-multiple")
async def upload_multiple_files(
    files: List[UploadFile] = File(...),
    user_id: UUID = Depends(get_current_user_id)
):
    """
    Upload multiple files at once.
    Returns extracted content for all files.
    """
    results = []
    
    for file in files:
        try:
            content_type = file.content_type or "application/octet-stream"
            
            if content_type not in SUPPORTED_TYPES:
                results.append({
                    "filename": file.filename,
                    "error": f"Unsupported file type: {content_type}"
                })
                continue
            
            file_bytes = await file.read()
            
            if len(file_bytes) > MAX_FILE_SIZE:
                results.append({
                    "filename": file.filename,
                    "error": f"File too large (max {MAX_FILE_SIZE / 1024 / 1024}MB)"
                })
                continue
            
            file_id = str(uuid.uuid4())
            filename = file.filename or f"upload_{file_id}"
            
            extracted = process_file(file_bytes, content_type, filename)
            
            results.append({
                "id": file_id,
                "filename": filename,
                "content_type": content_type,
                "size": len(file_bytes),
                "extracted_text": extracted.text if not extracted.metadata.get("is_image") else None,
                "data_uri": extracted.metadata.get("data_uri"),
                "metadata": extracted.metadata
            })
            
        except Exception as e:
            logger.error(f"Error processing file {file.filename}: {e}")
            results.append({
                "filename": file.filename,
                "error": str(e)
            })
    
    return {"files": results, "count": len(results)}


@router.post("/extract-text")
async def extract_text_from_data_uri(
    data_uri: str = Form(...),
    filename: str = Form(None),
    user_id: UUID = Depends(get_current_user_id)
):
    """
    Extract text from a base64 data URI.
    Useful for processing files already in the browser.
    """
    try:
        # Parse data URI
        if not data_uri.startswith("data:"):
            raise HTTPException(status_code=400, detail="Invalid data URI format")
        
        # Format: data:mime/type;base64,content
        header, content = data_uri.split(",", 1)
        mime_type = header.split(":")[1].split(";")[0]
        
        # Decode content
        file_bytes = base64.b64decode(content)
        
        if mime_type not in SUPPORTED_TYPES:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {mime_type}"
            )
        
        extracted = process_file(file_bytes, mime_type, filename or "uploaded_file")
        
        return {
            "text": extracted.text,
            "tables": extracted.tables,
            "metadata": extracted.metadata
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error extracting from data URI: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing file: {str(e)}")


@router.get("/supported-types")
async def get_supported_types():
    """Get list of supported file types"""
    return {
        "types": SUPPORTED_TYPES,
        "max_size_mb": MAX_FILE_SIZE / 1024 / 1024,
        "categories": {
            "documents": ["pdf", "doc", "docx", "txt", "md"],
            "spreadsheets": ["csv", "xls", "xlsx"],
            "data": ["json"],
            "images": ["png", "jpg", "gif", "webp"]
        }
    }
